"""PDF and DOCX exporters for completed deep-research reports."""

from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path

_MARKDOWN_MARKS = re.compile(r"(```|`|\*\*|__|\*|_)")
_LINK = re.compile(r"\[([^]]+)]\(([^)]+)\)")


def report_to_docx(markdown: str, *, title: str) -> bytes:
    from docx import Document
    from docx.shared import Inches

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    document.add_heading(title, level=0)
    for raw in markdown.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("### "):
            document.add_heading(_plain(line[4:]), level=3)
        elif line.startswith("## "):
            document.add_heading(_plain(line[3:]), level=2)
        elif line.startswith("# "):
            document.add_heading(_plain(line[2:]), level=1)
        elif line.startswith(("- ", "* ")):
            document.add_paragraph(_plain(line[2:]), style="List Bullet")
        elif re.match(r"^\d+[.)]\s+", line):
            document.add_paragraph(_plain(re.sub(r"^\d+[.)]\s+", "", line)), style="List Number")
        else:
            document.add_paragraph(_plain(line))
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def report_to_pdf(markdown: str, *, title: str) -> bytes:
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    font_name = "Helvetica"
    for candidate in (
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    ):
        if candidate.is_file():
            font_name = "ResearchUnicode"
            if font_name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(font_name, str(candidate)))
            break

    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
        title=title,
    )
    styles = getSampleStyleSheet()
    body = ParagraphStyle("ResearchBody", parent=styles["BodyText"], fontName=font_name)
    heading = ParagraphStyle(
        "ResearchHeading", parent=styles["Heading2"], fontName=font_name, spaceBefore=8
    )
    title_style = ParagraphStyle(
        "ResearchTitle",
        parent=styles["Title"],
        fontName=font_name,
        alignment=TA_CENTER,
    )
    story = [Paragraph(_xml(title), title_style), Spacer(1, 6 * mm)]
    for raw in markdown.splitlines():
        line = raw.strip()
        if not line:
            story.append(Spacer(1, 2 * mm))
            continue
        if line.startswith("#"):
            line = line.lstrip("# ")
            story.append(Paragraph(_xml(_plain(line)), heading))
        else:
            story.append(Paragraph(_xml(_plain(line)), body))
    document.build(story)
    return buffer.getvalue()


def _plain(text: str) -> str:
    text = _LINK.sub(r"\1 (\2)", text)
    return _MARKDOWN_MARKS.sub("", text)


def _xml(text: str) -> str:
    import html

    return html.escape(text).replace("\n", "<br/>")
